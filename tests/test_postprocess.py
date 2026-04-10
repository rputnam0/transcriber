from __future__ import annotations

import json
import sys
import types
from pathlib import Path

import pytest
from docx import Document

from transcriber import postprocess as postprocess_mod


def _write_prompt_files(prompts_dir: Path) -> None:
    prompt_files = {
        "analysis_system_prompt.txt": "You are the analyst.",
        "analysis_prompt_template.txt": (
            "BG={avarias_bg}\nOVW={campaign_ovw}\nTRANS={session_trans}\n"
            "INST={session_analysis_inst}\nEX={analysis_example}"
        ),
        "avarias_background.txt": "Background",
        "campaign_overview.txt": "Base overview",
        "session_analysis_instructions.txt": "Analyze it",
        "analysis_example.txt": "Example",
        "co_prompt_template.txt": (
            "PREV={previous_campaign_overview}\nANALYSIS={session_analysis}"
        ),
        "summary_system_prompt.txt": "You are the summarizer.",
        "summary_prompt_template.txt": (
            "BG={avarias_bg}\nOVW={campaign_ovw}\nTRANS={session_trans}\n"
            "ANALYSIS={session_analysis}\nINST={summary_inst}"
        ),
        "summary_instructions.txt": "Summarize it",
    }
    prompts_dir.mkdir(parents=True, exist_ok=True)
    for filename, content in prompt_files.items():
        (prompts_dir / filename).write_text(content, encoding="utf-8")


def _write_transcription_marker(
    transcript_path: Path,
    status: str,
    **extra_fields: object,
) -> None:
    marker_path = transcript_path.parent / f"{transcript_path.stem}.transcribe.json"
    payload = {"status": status, **extra_fields}
    marker_path.write_text(json.dumps(payload), encoding="utf-8")


def test_run_postprocess_for_transcript_creates_outputs(monkeypatch, tmp_path):
    prompts_dir = tmp_path / "prompts"
    summaries_dir = tmp_path / "summaries"
    transcript_dir = tmp_path / "transcripts" / "Session 32"
    transcript_dir.mkdir(parents=True, exist_ok=True)
    transcript_path = transcript_dir / "Session 32.txt"
    transcript_path.write_text("Speaker 00 00:00:00 Hello there\n", encoding="utf-8")
    _write_prompt_files(prompts_dir)

    config = postprocess_mod.PostProcessConfig(
        enabled=True,
        provider="google",
        model="test-model",
        prompts_dir=prompts_dir,
        summaries_dir=summaries_dir,
        api_key_env="GOOGLE_API_KEY",
        calls_per_minute=60,
    )

    class DummyGenerator:
        calls: list[tuple[str, str | None]] = []

        def __init__(self, cfg):
            assert cfg is config

        def generate(self, *, prompt: str, system_instruction: str | None = None) -> str:
            self.calls.append((prompt, system_instruction))
            if len(self.calls) == 1:
                return "# Analysis\n\nThe party solved a puzzle."
            if len(self.calls) == 2:
                return (
                    "<campaign_overview># Campaign Overview\n\nThings changed.</campaign_overview>"
                )
            return "<Summary># Summary\n\nSession recap.</Summary>"

    monkeypatch.setattr(postprocess_mod, "_GoogleTextGenerator", DummyGenerator)

    result = postprocess_mod.run_postprocess_for_transcript(transcript_path, config)

    assert result.created_marker is True
    assert result.paths.analysis_txt.exists()
    assert result.paths.analysis_docx.exists()
    assert result.paths.campaign_overview_txt.exists()
    assert result.paths.campaign_overview_docx.exists()
    assert result.paths.summary_txt.exists()
    assert result.paths.summary_docx.exists()
    assert result.paths.completion_marker.exists()

    marker = json.loads(result.paths.completion_marker.read_text(encoding="utf-8"))
    assert marker["status"] == "completed"
    assert marker["session_number"] == 32
    assert marker["provider"] == "google"

    summary_doc = Document(result.paths.summary_docx)
    joined_text = "\n".join(paragraph.text for paragraph in summary_doc.paragraphs)
    assert "Summary" in joined_text
    assert "Session recap." in joined_text

    second_result = postprocess_mod.run_postprocess_for_transcript(transcript_path, config)
    assert second_result.created_marker is False
    assert len(DummyGenerator.calls) == 3


def test_run_postprocess_for_transcript_regenerates_outputs_for_new_source(monkeypatch, tmp_path):
    prompts_dir = tmp_path / "prompts"
    summaries_dir = tmp_path / "summaries"
    transcript_dir = tmp_path / "transcripts" / "Session 65"
    transcript_dir.mkdir(parents=True, exist_ok=True)
    transcript_path = transcript_dir / "Session 65.txt"
    transcript_path.write_text("Speaker 00 00:00:00 Real session transcript\n", encoding="utf-8")
    _write_prompt_files(prompts_dir)

    config = postprocess_mod.PostProcessConfig(
        enabled=True,
        provider="google",
        model="test-model",
        prompts_dir=prompts_dir,
        summaries_dir=summaries_dir,
        api_key_env="GOOGLE_API_KEY",
        calls_per_minute=60,
    )

    paths = postprocess_mod.build_postprocess_paths(transcript_path, config)
    paths.raw_text_folder.mkdir(parents=True, exist_ok=True)
    stale_outputs = {
        paths.analysis_txt: "# Analysis\n\nStale analysis.",
        paths.campaign_overview_txt: "# Campaign Overview\n\nStale overview.",
        paths.summary_txt: "# Summary\n\nStale summary.",
    }
    for path, text in stale_outputs.items():
        path.write_text(text, encoding="utf-8")
    for path in (
        paths.analysis_docx,
        paths.campaign_overview_docx,
        paths.summary_docx,
    ):
        Document().save(path)
    paths.completion_marker.write_text(
        json.dumps(
            {
                "status": "completed",
                "transcript_path": str(
                    tmp_path / "transcripts" / "Session 65 test" / "Session 65 test.txt"
                ),
                "source_transcript_paths": [
                    str(tmp_path / "transcripts" / "Session 65 test" / "Session 65 test.txt")
                ],
            }
        ),
        encoding="utf-8",
    )

    class DummyGenerator:
        calls: list[tuple[str, str | None]] = []

        def __init__(self, cfg):
            assert cfg is config

        def generate(self, *, prompt: str, system_instruction: str | None = None) -> str:
            self.calls.append((prompt, system_instruction))
            if len(self.calls) == 1:
                return "# Analysis\n\nFresh analysis."
            if len(self.calls) == 2:
                return (
                    "<campaign_overview># Campaign Overview\n\nFresh overview.</campaign_overview>"
                )
            return "<Summary># Summary\n\nFresh summary.</Summary>"

    monkeypatch.setattr(postprocess_mod, "_GoogleTextGenerator", DummyGenerator)

    result = postprocess_mod.run_postprocess_for_transcript(transcript_path, config)

    assert result.created_marker is True
    assert len(DummyGenerator.calls) == 3
    assert paths.analysis_txt.read_text(encoding="utf-8") == "# Analysis\n\nFresh analysis.\n"
    assert (
        paths.campaign_overview_txt.read_text(encoding="utf-8")
        == "# Campaign Overview\n\nFresh overview.\n"
    )
    assert paths.summary_txt.read_text(encoding="utf-8") == "# Summary\n\nFresh summary.\n"

    marker = json.loads(paths.completion_marker.read_text(encoding="utf-8"))
    assert marker["transcript_path"] == str(transcript_path)
    assert marker["source_transcript_paths"] == [str(transcript_path)]


def test_run_postprocess_for_transcript_writes_summary_ready_receipt(monkeypatch, tmp_path):
    prompts_dir = tmp_path / "prompts"
    summaries_dir = tmp_path / "summaries"
    state_dir = tmp_path / "state"
    transcript_dir = tmp_path / "transcripts" / "Session 65"
    transcript_dir.mkdir(parents=True, exist_ok=True)
    transcript_path = transcript_dir / "Session 65.txt"
    transcript_path.write_text("Speaker 00 00:00:00 Real session transcript\n", encoding="utf-8")
    _write_prompt_files(prompts_dir)
    _write_transcription_marker(
        transcript_path,
        "completed",
        input_path=str(tmp_path / "incoming" / "Session 65.zip"),
        recording_id="rec-65",
        title="Session 65",
    )
    state_dir.mkdir(parents=True, exist_ok=True)
    (state_dir / "rec-65.drive-uploaded.json").write_text(
        json.dumps({"recording_id": "rec-65", "url": "https://drive.example/rec-65"}),
        encoding="utf-8",
    )

    config = postprocess_mod.PostProcessConfig(
        enabled=True,
        provider="google",
        model="gemini-3-flash-preview",
        prompts_dir=prompts_dir,
        summaries_dir=summaries_dir,
        api_key_env="GOOGLE_API_KEY",
        calls_per_minute=60,
        summary_ready=postprocess_mod.SummaryReadyConfig(
            state_target=str(state_dir),
            scenes_count=2,
        ),
    )

    class DummyGenerator:
        calls: list[tuple[str, str | None]] = []

        def __init__(self, cfg):
            assert cfg is config

        def generate(self, *, prompt: str, system_instruction: str | None = None) -> str:
            self.calls.append((prompt, system_instruction))
            if len(self.calls) == 1:
                return "# Analysis\n\nFresh analysis."
            if len(self.calls) == 2:
                return (
                    "<campaign_overview># Campaign Overview\n\nFresh overview.</campaign_overview>"
                )
            if len(self.calls) == 3:
                return "<Summary># Summary\n\nFresh summary.</Summary>"
            return (
                "<announcement_scenes>"
                "<scene>The bard turned a hostage negotiation into stand-up.</scene>"
                "<scene>The fighter kicked open the trapped door immediately after warning everyone.</scene>"
                "</announcement_scenes>"
            )

    monkeypatch.setattr(postprocess_mod, "_GoogleTextGenerator", DummyGenerator)
    monkeypatch.setattr(
        postprocess_mod,
        "_resolve_summary_document_url",
        lambda path: "https://drive.example/docs/session-65",
    )

    result = postprocess_mod.run_postprocess_for_transcript(transcript_path, config)

    assert result.created_marker is True
    receipt = json.loads((state_dir / "rec-65.summary-ready.json").read_text(encoding="utf-8"))
    assert receipt["recording_id"] == "rec-65"
    assert receipt["title"] == "Session 65"
    assert receipt["drive_url"] == "https://drive.example/rec-65"
    assert receipt["document_url"] == "https://drive.example/docs/session-65"
    assert receipt["announcement_scenes"] == [
        "The bard turned a hostage negotiation into stand-up.",
        "The fighter kicked open the trapped door immediately after warning everyone.",
    ]
    assert receipt["summary_excerpt"] == "# Summary Fresh summary."
    summary_ready_marker = result.paths.session_folder / "session_65.summary-ready.json"
    assert summary_ready_marker.exists()
    assert len(DummyGenerator.calls) == 4


def test_run_postprocess_for_transcript_backfills_drive_uploaded_receipt_from_audio_bundle(
    monkeypatch,
    tmp_path,
):
    prompts_dir = tmp_path / "prompts"
    summaries_dir = tmp_path / "summaries"
    state_dir = tmp_path / "state"
    audio_dir = tmp_path / "drive" / "Audio"
    incoming_dir = tmp_path / "incoming"
    transcript_dir = tmp_path / "transcripts" / "Session 66"
    transcript_dir.mkdir(parents=True, exist_ok=True)
    incoming_dir.mkdir(parents=True, exist_ok=True)
    transcript_path = transcript_dir / "Session 66.txt"
    source_zip = incoming_dir / "Session 66.zip"
    source_zip.write_bytes(b"zip-bytes-66")
    (incoming_dir / "Session 66.json").write_text(
        json.dumps(
            {
                "recording_id": "rec-66",
                "preferred_basename": "Session 66",
                "requester_id": "user-66",
                "format": "flac",
                "container": "zip",
            }
        ),
        encoding="utf-8",
    )
    transcript_path.write_text("Speaker 00 00:00:00 Real session transcript\n", encoding="utf-8")
    _write_prompt_files(prompts_dir)
    _write_transcription_marker(
        transcript_path,
        "completed",
        input_path=str(source_zip),
        recording_id="rec-66",
        title="Session 66",
    )
    state_dir.mkdir(parents=True, exist_ok=True)

    config = postprocess_mod.PostProcessConfig(
        enabled=True,
        provider="google",
        model="gemini-3-flash-preview",
        prompts_dir=prompts_dir,
        summaries_dir=summaries_dir,
        api_key_env="GOOGLE_API_KEY",
        calls_per_minute=60,
        summary_ready=postprocess_mod.SummaryReadyConfig(
            state_target=str(state_dir),
            scenes_count=2,
            audio_dir=audio_dir,
            audio_link_wait_seconds=0,
        ),
    )

    class DummyGenerator:
        calls: list[tuple[str, str | None]] = []

        def __init__(self, cfg):
            assert cfg is config

        def generate(self, *, prompt: str, system_instruction: str | None = None) -> str:
            self.calls.append((prompt, system_instruction))
            if len(self.calls) == 1:
                return "# Analysis\n\nFresh analysis."
            if len(self.calls) == 2:
                return (
                    "<campaign_overview># Campaign Overview\n\nFresh overview.</campaign_overview>"
                )
            if len(self.calls) == 3:
                return "<Summary># Summary\n\nFresh summary.</Summary>"
            return (
                "<announcement_scenes>"
                "<scene>The goblin king dissolved into pure courtroom chaos.</scene>"
                "<scene>The rogue pickpocketed the one guard already on their side.</scene>"
                "</announcement_scenes>"
            )

    captured: dict[str, Path] = {}

    def fake_wait_for_drivefs_cloud_id(path: Path, *, timeout_seconds: int) -> str | None:
        captured["audio_path"] = path
        assert timeout_seconds == 0
        assert path.exists()
        return "audio-file-66"

    monkeypatch.setattr(postprocess_mod, "_GoogleTextGenerator", DummyGenerator)
    monkeypatch.setattr(
        postprocess_mod,
        "_resolve_summary_document_url",
        lambda path: "https://drive.example/docs/session-66",
    )
    monkeypatch.setattr(
        postprocess_mod,
        "_wait_for_drivefs_cloud_id",
        fake_wait_for_drivefs_cloud_id,
    )

    result = postprocess_mod.run_postprocess_for_transcript(transcript_path, config)

    assert result.created_marker is True
    copied_audio = audio_dir / "Session 66.zip"
    assert copied_audio.read_bytes() == b"zip-bytes-66"
    assert captured["audio_path"] == copied_audio

    drive_receipt = json.loads(
        (state_dir / "rec-66.drive-uploaded.json").read_text(encoding="utf-8")
    )
    assert drive_receipt["recording_id"] == "rec-66"
    assert drive_receipt["user_id"] == "user-66"
    assert drive_receipt["file_id"] == "audio-file-66"
    assert drive_receipt["url"] == "https://drive.google.com/open?id=audio-file-66"
    assert drive_receipt["format"] == "flac"
    assert drive_receipt["container"] == "zip"

    summary_receipt = json.loads(
        (state_dir / "rec-66.summary-ready.json").read_text(encoding="utf-8")
    )
    assert summary_receipt["drive_url"] == "https://drive.google.com/open?id=audio-file-66"
    assert summary_receipt["document_url"] == "https://drive.example/docs/session-66"


def test_run_postprocess_for_transcript_writes_summary_ready_receipt_without_drive_upload(
    monkeypatch,
    tmp_path,
):
    prompts_dir = tmp_path / "prompts"
    summaries_dir = tmp_path / "summaries"
    state_dir = tmp_path / "state"
    transcript_dir = tmp_path / "transcripts" / "Session 66"
    transcript_dir.mkdir(parents=True, exist_ok=True)
    transcript_path = transcript_dir / "Session 66.txt"
    transcript_path.write_text("Speaker 00 00:00:00 Real session transcript\n", encoding="utf-8")
    _write_prompt_files(prompts_dir)
    _write_transcription_marker(
        transcript_path,
        "completed",
        input_path=str(tmp_path / "incoming" / "Session 66.zip"),
        recording_id="rec-66",
        title="Session 66",
    )
    state_dir.mkdir(parents=True, exist_ok=True)

    config = postprocess_mod.PostProcessConfig(
        enabled=True,
        provider="google",
        model="gemini-3-flash-preview",
        prompts_dir=prompts_dir,
        summaries_dir=summaries_dir,
        api_key_env="GOOGLE_API_KEY",
        calls_per_minute=60,
        summary_ready=postprocess_mod.SummaryReadyConfig(
            state_target=str(state_dir),
            scenes_count=2,
        ),
    )

    class DummyGenerator:
        calls: list[tuple[str, str | None]] = []

        def __init__(self, cfg):
            assert cfg is config

        def generate(self, *, prompt: str, system_instruction: str | None = None) -> str:
            self.calls.append((prompt, system_instruction))
            if len(self.calls) == 1:
                return "# Analysis\n\nFresh analysis."
            if len(self.calls) == 2:
                return (
                    "<campaign_overview># Campaign Overview\n\nFresh overview.</campaign_overview>"
                )
            if len(self.calls) == 3:
                return "<Summary># Summary\n\nFresh summary.</Summary>"
            return (
                "<announcement_scenes>"
                "<scene>The goblin king dissolved into pure courtroom chaos.</scene>"
                "<scene>The rogue pickpocketed the one guard already on their side.</scene>"
                "</announcement_scenes>"
            )

    monkeypatch.setattr(postprocess_mod, "_GoogleTextGenerator", DummyGenerator)
    monkeypatch.setattr(
        postprocess_mod,
        "_resolve_summary_document_url",
        lambda path: "https://drive.example/docs/session-66",
    )

    result = postprocess_mod.run_postprocess_for_transcript(transcript_path, config)

    assert result.created_marker is True
    receipt = json.loads((state_dir / "rec-66.summary-ready.json").read_text(encoding="utf-8"))
    assert receipt["recording_id"] == "rec-66"
    assert receipt["title"] == "Session 66"
    assert receipt["drive_url"] == ""
    assert receipt["document_url"] == "https://drive.example/docs/session-66"
    assert receipt["announcement_scenes"] == [
        "The goblin king dissolved into pure courtroom chaos.",
        "The rogue pickpocketed the one guard already on their side.",
    ]
    summary_ready_marker = result.paths.session_folder / "session_66.summary-ready.json"
    assert summary_ready_marker.exists()
    assert len(DummyGenerator.calls) == 4


def test_run_postprocess_for_transcript_retries_summary_receipt_without_regenerating_outputs(
    monkeypatch,
    tmp_path,
):
    prompts_dir = tmp_path / "prompts"
    summaries_dir = tmp_path / "summaries"
    state_dir = tmp_path / "state"
    transcript_dir = tmp_path / "transcripts" / "Session 65"
    transcript_dir.mkdir(parents=True, exist_ok=True)
    transcript_path = transcript_dir / "Session 65.txt"
    transcript_path.write_text("Speaker 00 00:00:00 Real session transcript\n", encoding="utf-8")
    _write_prompt_files(prompts_dir)
    _write_transcription_marker(
        transcript_path,
        "completed",
        input_path=str(tmp_path / "incoming" / "Session 65.zip"),
        recording_id="rec-65",
        title="Session 65",
    )
    state_dir.mkdir(parents=True, exist_ok=True)
    (state_dir / "rec-65.drive-uploaded.json").write_text(
        json.dumps({"recording_id": "rec-65", "url": "https://drive.example/rec-65"}),
        encoding="utf-8",
    )

    config = postprocess_mod.PostProcessConfig(
        enabled=True,
        provider="google",
        model="gemini-3-flash-preview",
        prompts_dir=prompts_dir,
        summaries_dir=summaries_dir,
        api_key_env="GOOGLE_API_KEY",
        calls_per_minute=60,
        summary_ready=postprocess_mod.SummaryReadyConfig(
            state_target=str(state_dir),
            scenes_count=2,
        ),
    )
    paths = postprocess_mod.build_postprocess_paths(transcript_path, config)
    paths.raw_text_folder.mkdir(parents=True, exist_ok=True)
    paths.analysis_txt.write_text("# Analysis\n\nFresh analysis.\n", encoding="utf-8")
    paths.campaign_overview_txt.write_text(
        "# Campaign Overview\n\nFresh overview.\n",
        encoding="utf-8",
    )
    paths.summary_txt.write_text("# Summary\n\nFresh summary.\n", encoding="utf-8")
    for path in (
        paths.analysis_docx,
        paths.campaign_overview_docx,
        paths.summary_docx,
    ):
        Document().save(path)
    paths.completion_marker.write_text(
        json.dumps(
            {
                "status": "completed",
                "transcript_path": str(transcript_path),
                "source_transcript_paths": [str(transcript_path)],
            }
        ),
        encoding="utf-8",
    )

    class DummyGenerator:
        calls: list[tuple[str, str | None]] = []

        def __init__(self, cfg):
            assert cfg is config

        def generate(self, *, prompt: str, system_instruction: str | None = None) -> str:
            self.calls.append((prompt, system_instruction))
            return (
                "<announcement_scenes>"
                "<scene>The wizard cast a speech instead of a spell.</scene>"
                "<scene>The cleric survived entirely on indignation.</scene>"
                "</announcement_scenes>"
            )

    monkeypatch.setattr(postprocess_mod, "_GoogleTextGenerator", DummyGenerator)
    monkeypatch.setattr(
        postprocess_mod,
        "_resolve_summary_document_url",
        lambda path: "https://drive.example/docs/session-65",
    )

    result = postprocess_mod.run_postprocess_for_transcript(transcript_path, config)

    assert result.created_marker is False
    assert len(DummyGenerator.calls) == 1
    receipt = json.loads((state_dir / "rec-65.summary-ready.json").read_text(encoding="utf-8"))
    assert receipt["announcement_scenes"] == [
        "The wizard cast a speech instead of a spell.",
        "The cleric survived entirely on indignation.",
    ]
    assert receipt["document_url"] == "https://drive.example/docs/session-65"


def test_run_postprocess_for_transcript_refreshes_summary_ready_when_drive_receipt_arrives(
    monkeypatch,
    tmp_path,
):
    prompts_dir = tmp_path / "prompts"
    summaries_dir = tmp_path / "summaries"
    state_dir = tmp_path / "state"
    transcript_dir = tmp_path / "transcripts" / "Session 65"
    transcript_dir.mkdir(parents=True, exist_ok=True)
    transcript_path = transcript_dir / "Session 65.txt"
    transcript_path.write_text("Speaker 00 00:00:00 Real session transcript\n", encoding="utf-8")
    _write_prompt_files(prompts_dir)
    _write_transcription_marker(
        transcript_path,
        "completed",
        input_path=str(tmp_path / "incoming" / "Session 65.zip"),
        recording_id="rec-65",
        title="Session 65",
    )
    state_dir.mkdir(parents=True, exist_ok=True)
    (state_dir / "rec-65.drive-uploaded.json").write_text(
        json.dumps({"recording_id": "rec-65", "url": "https://drive.example/audio/session-65"}),
        encoding="utf-8",
    )

    config = postprocess_mod.PostProcessConfig(
        enabled=True,
        provider="google",
        model="gemini-3-flash-preview",
        prompts_dir=prompts_dir,
        summaries_dir=summaries_dir,
        api_key_env="GOOGLE_API_KEY",
        calls_per_minute=60,
        summary_ready=postprocess_mod.SummaryReadyConfig(
            state_target=str(state_dir),
            scenes_count=2,
        ),
    )
    paths = postprocess_mod.build_postprocess_paths(transcript_path, config)
    paths.raw_text_folder.mkdir(parents=True, exist_ok=True)
    paths.analysis_txt.write_text("# Analysis\n\nFresh analysis.\n", encoding="utf-8")
    paths.campaign_overview_txt.write_text(
        "# Campaign Overview\n\nFresh overview.\n",
        encoding="utf-8",
    )
    paths.summary_txt.write_text("# Summary\n\nFresh summary.\n", encoding="utf-8")
    for path in (
        paths.analysis_docx,
        paths.campaign_overview_docx,
        paths.summary_docx,
    ):
        Document().save(path)
    paths.completion_marker.write_text(
        json.dumps(
            {
                "status": "completed",
                "transcript_path": str(transcript_path),
                "source_transcript_paths": [str(transcript_path)],
            }
        ),
        encoding="utf-8",
    )
    (paths.session_folder / "session_65.summary-ready.json").write_text(
        json.dumps(
            {
                "recording_id": "rec-65",
                "transcript_path": str(transcript_path),
                "source_transcript_paths": [str(transcript_path)],
                "summary_path": str(paths.summary_txt),
                "drive_url": "",
                "document_url": "https://drive.example/docs/session-65",
                "published_at": "2026-04-10T06:00:00Z",
            }
        ),
        encoding="utf-8",
    )

    class DummyGenerator:
        calls: list[tuple[str, str | None]] = []

        def __init__(self, cfg):
            assert cfg is config

        def generate(self, *, prompt: str, system_instruction: str | None = None) -> str:
            self.calls.append((prompt, system_instruction))
            return (
                "<announcement_scenes>"
                "<scene>The wizard cast a speech instead of a spell.</scene>"
                "<scene>The cleric survived entirely on indignation.</scene>"
                "</announcement_scenes>"
            )

    monkeypatch.setattr(postprocess_mod, "_GoogleTextGenerator", DummyGenerator)
    monkeypatch.setattr(
        postprocess_mod,
        "_resolve_summary_document_url",
        lambda path: "https://drive.example/docs/session-65",
    )

    result = postprocess_mod.run_postprocess_for_transcript(transcript_path, config)

    assert result.created_marker is False
    assert len(DummyGenerator.calls) == 1
    receipt = json.loads((state_dir / "rec-65.summary-ready.json").read_text(encoding="utf-8"))
    assert receipt["drive_url"] == "https://drive.example/audio/session-65"
    marker = json.loads(
        (paths.session_folder / "session_65.summary-ready.json").read_text(encoding="utf-8")
    )
    assert marker["drive_url"] == "https://drive.example/audio/session-65"


def test_resolve_summary_document_url_from_drivefs_db(monkeypatch, tmp_path):
    db_path = tmp_path / "mirror_metadata_sqlite.db"
    conn = postprocess_mod.sqlite3.connect(str(db_path))
    try:
        conn.execute(
            """
            create table items (
                stable_id integer primary key,
                id text,
                proto blob,
                trashed boolean,
                starred boolean,
                is_owner boolean,
                mime_type text,
                is_folder boolean,
                modified_date integer,
                shared_with_me_date integer,
                viewed_by_me_date integer,
                file_size integer,
                is_tombstone boolean,
                local_title text,
                subscribed boolean,
                team_drive_stable_id integer,
                local_title_tokenized text,
                inaccessible_inheritance_broken boolean
            )
            """
        )
        conn.execute(
            """
            create table stable_ids (
                stable_id integer primary key,
                cloud_id text
            )
            """
        )
        conn.execute(
            """
            create table stable_parents (
                item_stable_id integer,
                parent_stable_id integer,
                local_title_hash integer,
                primary key (item_stable_id, parent_stable_id)
            )
            """
        )
        rows = [
            (101, None, False, "My Drive"),
            (219, None, False, "DND"),
            (218, None, False, "Summaries"),
            (30844, None, False, "Session 65"),
            (30845, None, False, "session_65_summary_output.docx"),
            (274, None, False, "Junk Drawer"),
            (30824, None, False, "Session 65"),
            (30827, None, False, "session_65_summary_output.docx"),
        ]
        for stable_id, cloud_id, is_folder, local_title in rows:
            conn.execute(
                """
                insert into items (
                    stable_id, id, proto, trashed, starred, is_owner, mime_type, is_folder,
                    modified_date, shared_with_me_date, viewed_by_me_date, file_size, is_tombstone,
                    local_title, subscribed, team_drive_stable_id, local_title_tokenized,
                    inaccessible_inheritance_broken
                ) values (?, ?, null, 0, 0, 1, '', ?, 0, 0, 0, 0, 0, ?, 1, null, ?, 0)
                """,
                (stable_id, cloud_id or "", int(is_folder), local_title, local_title.lower()),
            )
        conn.executemany(
            "insert into stable_ids (stable_id, cloud_id) values (?, ?)",
            [
                (30845, "doc-good"),
                (30827, "doc-stale"),
            ],
        )
        conn.executemany(
            "insert into stable_parents (item_stable_id, parent_stable_id, local_title_hash) values (?, ?, 0)",
            [
                (219, 101),
                (218, 219),
                (30844, 218),
                (30845, 30844),
                (274, 101),
                (30824, 274),
                (30827, 30824),
            ],
        )
        conn.commit()
    finally:
        conn.close()

    monkeypatch.setattr(postprocess_mod, "_drivefs_metadata_db_paths", lambda: [db_path])

    url = postprocess_mod._resolve_summary_document_url(
        Path("/mnt/g/My Drive/DND/Summaries/Session 65/session_65_summary_output.docx")
    )

    assert url == "https://drive.google.com/open?id=doc-good"


def test_resolve_summary_document_url_reads_recent_drivefs_entries_from_wal(monkeypatch, tmp_path):
    db_path = tmp_path / "mirror_metadata_sqlite.db"
    conn = postprocess_mod.sqlite3.connect(str(db_path))
    try:
        conn.execute("pragma journal_mode=wal")
        conn.execute(
            """
            create table items (
                stable_id integer primary key,
                id text,
                proto blob,
                trashed boolean,
                starred boolean,
                is_owner boolean,
                mime_type text,
                is_folder boolean,
                modified_date integer,
                shared_with_me_date integer,
                viewed_by_me_date integer,
                file_size integer,
                is_tombstone boolean,
                local_title text,
                subscribed boolean,
                team_drive_stable_id integer,
                local_title_tokenized text,
                inaccessible_inheritance_broken boolean
            )
            """
        )
        conn.execute(
            """
            create table stable_ids (
                stable_id integer primary key,
                cloud_id text
            )
            """
        )
        conn.execute(
            """
            create table stable_parents (
                item_stable_id integer,
                parent_stable_id integer,
                local_title_hash integer,
                primary key (item_stable_id, parent_stable_id)
            )
            """
        )
        rows = [
            (101, None, True, "My Drive"),
            (219, None, True, "DND"),
            (218, None, True, "Summaries"),
            (30866, None, True, "Session 66"),
            (30867, None, False, "session_66_summary_output.docx"),
        ]
        for stable_id, cloud_id, is_folder, local_title in rows:
            conn.execute(
                """
                insert into items (
                    stable_id, id, proto, trashed, starred, is_owner, mime_type, is_folder,
                    modified_date, shared_with_me_date, viewed_by_me_date, file_size, is_tombstone,
                    local_title, subscribed, team_drive_stable_id, local_title_tokenized,
                    inaccessible_inheritance_broken
                ) values (?, ?, null, 0, 0, 1, '', ?, 0, 0, 0, 0, 0, ?, 1, null, ?, 0)
                """,
                (stable_id, cloud_id or "", int(is_folder), local_title, local_title.lower()),
            )
        conn.execute(
            "insert into stable_ids (stable_id, cloud_id) values (?, ?)",
            (30867, "doc-from-wal"),
        )
        conn.executemany(
            "insert into stable_parents (item_stable_id, parent_stable_id, local_title_hash) values (?, ?, 0)",
            [
                (219, 101),
                (218, 219),
                (30866, 218),
                (30867, 30866),
            ],
        )
        conn.commit()

        monkeypatch.setattr(postprocess_mod, "_drivefs_metadata_db_paths", lambda: [db_path])

        url = postprocess_mod._resolve_summary_document_url(
            Path("/mnt/g/My Drive/DND/Summaries/Session 66/session_66_summary_output.docx")
        )
    finally:
        conn.close()

    assert url == "https://drive.google.com/open?id=doc-from-wal"


def test_expected_completion_marker_path_uses_session_folder(tmp_path):
    config = postprocess_mod.PostProcessConfig(
        enabled=True,
        provider="google",
        model="test-model",
        prompts_dir=tmp_path / "prompts",
        summaries_dir=tmp_path / "summaries",
    )

    marker = postprocess_mod.expected_completion_marker_path(
        tmp_path / "transcripts" / "Session 7" / "Session 7.txt",
        config,
    )

    assert marker == tmp_path / "summaries" / "Session 7" / "session_7.postprocess.json"


def test_postprocess_delay_seconds_allows_no_throttle(tmp_path):
    config = postprocess_mod.PostProcessConfig(
        enabled=True,
        provider="google",
        model="test-model",
        prompts_dir=tmp_path / "prompts",
        summaries_dir=tmp_path / "summaries",
        calls_per_minute=0,
    )

    assert config.delay_seconds == 0.0


def test_run_postprocess_for_non_session_transcript_raises_clear_error(tmp_path):
    prompts_dir = tmp_path / "prompts"
    summaries_dir = tmp_path / "summaries"
    transcript_dir = tmp_path / "transcripts" / "Live Smoke Test"
    transcript_dir.mkdir(parents=True, exist_ok=True)
    transcript_path = transcript_dir / "Live Smoke Test.txt"
    transcript_path.write_text("Speaker 00 00:00:00 Hello there\n", encoding="utf-8")
    _write_prompt_files(prompts_dir)

    config = postprocess_mod.PostProcessConfig(
        enabled=True,
        provider="google",
        model="test-model",
        prompts_dir=prompts_dir,
        summaries_dir=summaries_dir,
    )

    with pytest.raises(postprocess_mod.UnsupportedTranscriptPostprocessError):
        postprocess_mod.run_postprocess_for_transcript(transcript_path, config)


def test_can_postprocess_transcript_requires_explicit_session_name() -> None:
    assert postprocess_mod.can_postprocess_transcript("Session 3.txt")
    assert postprocess_mod.can_postprocess_transcript("Session 28 1_2.txt")
    assert not postprocess_mod.can_postprocess_transcript("kZWVHz3eCBaf.txt")
    assert not postprocess_mod.can_postprocess_transcript(
        "Odds and Ends Transcripts/kZWVHz3eCBaf.txt"
    )


def test_resolve_postprocess_config_accepts_thinking_level(tmp_path):
    cfg = {
        "postprocess": {
            "enabled": True,
            "provider": "google",
            "model": "gemini-3-flash-preview",
            "thinking_level": "high",
            "prompts_dir": str(tmp_path / "prompts"),
            "summaries_dir": str(tmp_path / "summaries"),
        }
    }

    config = postprocess_mod.resolve_postprocess_config(cfg)

    assert config is not None
    assert config.model == "gemini-3-flash-preview"
    assert config.thinking_level == "high"


def test_resolve_postprocess_config_accepts_openai_provider(tmp_path):
    cfg = {
        "postprocess": {
            "enabled": True,
            "provider": "openai",
            "model": "gemma-4-26B-A4B-it-UD-Q3_K_M.gguf",
            "api_key_env": "LOCAL_LLM_API_KEY",
            "api_base_env": "LOCAL_LLM_API_BASE",
            "max_output_tokens": 768,
            "prompts_dir": str(tmp_path / "prompts"),
            "summaries_dir": str(tmp_path / "summaries"),
        }
    }

    config = postprocess_mod.resolve_postprocess_config(cfg)

    assert config is not None
    assert config.provider == "openai"
    assert config.api_key_env == "LOCAL_LLM_API_KEY"
    assert config.api_base_env == "LOCAL_LLM_API_BASE"
    assert config.max_output_tokens == 768


def test_resolve_postprocess_config_accepts_summary_audio_dir(tmp_path):
    cfg = {
        "postprocess": {
            "enabled": True,
            "provider": "google",
            "model": "gemini-3-flash-preview",
            "prompts_dir": str(tmp_path / "prompts"),
            "summaries_dir": str(tmp_path / "summaries"),
            "summary_ready": {
                "enabled": True,
                "state_target": "raspi:/srv/craig-stack/exports/state",
                "audio_dir": str(tmp_path / "drive" / "Audio"),
                "audio_link_wait_seconds": 15,
            },
        }
    }

    config = postprocess_mod.resolve_postprocess_config(cfg)

    assert config is not None
    assert config.summary_ready is not None
    assert config.summary_ready.audio_dir == (tmp_path / "drive" / "Audio")
    assert config.summary_ready.audio_link_wait_seconds == 15


def test_google_text_generator_passes_thinking_level_for_gemini_3(monkeypatch, tmp_path):
    captured: dict[str, object] = {}

    class FakeThinkingConfig:
        def __init__(self, **kwargs):
            captured["thinking_config_kwargs"] = kwargs
            self.kwargs = kwargs

    class FakeGenerateContentConfig:
        def __init__(self, **kwargs):
            captured["generate_config_kwargs"] = kwargs
            self.kwargs = kwargs

    class FakePart:
        @staticmethod
        def from_text(*, text: str):
            return {"text": text}

    class FakeContent:
        def __init__(self, *, role: str, parts: list[object]):
            self.role = role
            self.parts = parts

    class FakeModels:
        def generate_content(self, *, model, contents, config):
            captured["model"] = model
            captured["contents"] = contents
            captured["config"] = config
            return types.SimpleNamespace(text="ok")

    class FakeClient:
        def __init__(self, *, api_key: str):
            captured["api_key"] = api_key
            self.models = FakeModels()

    fake_types = types.SimpleNamespace(
        ThinkingConfig=FakeThinkingConfig,
        GenerateContentConfig=FakeGenerateContentConfig,
        Content=FakeContent,
        Part=FakePart,
    )
    fake_genai = types.SimpleNamespace(Client=FakeClient, types=fake_types)

    monkeypatch.setenv("GOOGLE_API_KEY", "test-key")
    monkeypatch.setitem(sys.modules, "google", types.SimpleNamespace(genai=fake_genai))
    monkeypatch.setitem(sys.modules, "google.genai", fake_genai)
    monkeypatch.setitem(sys.modules, "google.genai.types", fake_types)

    config = postprocess_mod.PostProcessConfig(
        enabled=True,
        provider="google",
        model="gemini-3-flash-preview",
        prompts_dir=tmp_path / "prompts",
        summaries_dir=tmp_path / "summaries",
        thinking_level="high",
    )

    generator = postprocess_mod._GoogleTextGenerator(config)
    text = generator.generate(prompt="hello", system_instruction="system")

    assert text == "ok"
    assert captured["api_key"] == "test-key"
    assert captured["model"] == "gemini-3-flash-preview"
    assert captured["thinking_config_kwargs"] == {"thinking_level": "high"}
    assert (
        captured["generate_config_kwargs"]["thinking_config"]
        is captured["config"].kwargs["thinking_config"]
    )


def test_openai_text_generator_uses_reasoning_content_when_content_is_empty(monkeypatch, tmp_path):
    captured: dict[str, object] = {}

    class FakeResponse:
        def __enter__(self):
            return self

        def __exit__(self, exc_type, exc, tb):
            return False

        def read(self) -> bytes:
            return json.dumps(
                {
                    "choices": [
                        {
                            "message": {
                                "content": "",
                                "reasoning_content": "<Summary># Summary\n\nGemma fallback summary.</Summary>",
                            }
                        }
                    ]
                }
            ).encode("utf-8")

    def fake_urlopen(req, timeout=0):
        captured["timeout"] = timeout
        captured["url"] = req.full_url
        captured["headers"] = dict(req.header_items())
        captured["body"] = json.loads(req.data.decode("utf-8"))
        return FakeResponse()

    monkeypatch.setenv("LOCAL_LLM_API_KEY", "dummy")
    monkeypatch.setenv("LOCAL_LLM_API_BASE", "http://127.0.0.1:8082/v1")
    monkeypatch.setattr(postprocess_mod.request, "urlopen", fake_urlopen)

    config = postprocess_mod.PostProcessConfig(
        enabled=True,
        provider="openai",
        model="gemma-4-26B-A4B-it-UD-Q3_K_M.gguf",
        prompts_dir=tmp_path / "prompts",
        summaries_dir=tmp_path / "summaries",
        api_key_env="LOCAL_LLM_API_KEY",
        api_base_env="LOCAL_LLM_API_BASE",
        max_output_tokens=512,
    )

    generator = postprocess_mod._OpenAITextGenerator(config)
    text = generator.generate(prompt="hello", system_instruction="system")

    assert text == "<Summary># Summary\n\nGemma fallback summary.</Summary>"
    assert captured["url"] == "http://127.0.0.1:8082/v1/chat/completions"
    assert captured["body"]["max_tokens"] == 512
    assert captured["body"]["messages"][0] == {"role": "system", "content": "system"}
    assert captured["body"]["messages"][1] == {"role": "user", "content": "hello"}


def test_split_session_ready_for_postprocess_requires_all_parts(tmp_path):
    transcript_path = tmp_path / "transcripts" / "Session 28 1_2" / "Session 28 1_2.txt"
    transcript_path.parent.mkdir(parents=True, exist_ok=True)
    transcript_path.write_text("part one", encoding="utf-8")

    assert postprocess_mod.split_session_ready_for_postprocess(transcript_path) is False


def test_split_session_ready_for_postprocess_requires_completed_part_markers(tmp_path):
    transcript_root = tmp_path / "transcripts"
    part_one = transcript_root / "Session 28 1_2" / "Session 28 1_2.txt"
    part_two = transcript_root / "Session 28 2_2" / "Session 28 2_2.txt"
    part_one.parent.mkdir(parents=True, exist_ok=True)
    part_two.parent.mkdir(parents=True, exist_ok=True)
    part_one.write_text("part one", encoding="utf-8")
    part_two.write_text("part two", encoding="utf-8")
    _write_transcription_marker(part_one, "completed")
    _write_transcription_marker(part_two, "in_progress")

    assert postprocess_mod.split_session_ready_for_postprocess(part_one) is False


def test_run_postprocess_for_split_transcript_raises_until_all_parts_exist(monkeypatch, tmp_path):
    prompts_dir = tmp_path / "prompts"
    summaries_dir = tmp_path / "summaries"
    transcript_path = tmp_path / "transcripts" / "Session 28 1_2" / "Session 28 1_2.txt"
    transcript_path.parent.mkdir(parents=True, exist_ok=True)
    transcript_path.write_text("part one", encoding="utf-8")
    _write_prompt_files(prompts_dir)

    config = postprocess_mod.PostProcessConfig(
        enabled=True,
        provider="google",
        model="test-model",
        prompts_dir=prompts_dir,
        summaries_dir=summaries_dir,
    )

    class DummyGenerator:
        def __init__(self, _cfg):
            pytest.fail("Generator should not run while split-session parts are missing.")

    monkeypatch.setattr(postprocess_mod, "_GoogleTextGenerator", DummyGenerator)

    with pytest.raises(postprocess_mod.SplitSessionPendingError):
        postprocess_mod.run_postprocess_for_transcript(transcript_path, config)


def test_run_postprocess_for_split_transcript_waits_for_completed_part_markers(
    monkeypatch,
    tmp_path,
):
    prompts_dir = tmp_path / "prompts"
    summaries_dir = tmp_path / "summaries"
    transcript_root = tmp_path / "transcripts"
    part_one = transcript_root / "Session 28 1_2" / "Session 28 1_2.txt"
    part_two = transcript_root / "Session 28 2_2" / "Session 28 2_2.txt"
    part_one.parent.mkdir(parents=True, exist_ok=True)
    part_two.parent.mkdir(parents=True, exist_ok=True)
    part_one.write_text("part one", encoding="utf-8")
    part_two.write_text("part two", encoding="utf-8")
    _write_prompt_files(prompts_dir)
    _write_transcription_marker(part_one, "completed")
    _write_transcription_marker(part_two, "in_progress")

    config = postprocess_mod.PostProcessConfig(
        enabled=True,
        provider="google",
        model="test-model",
        prompts_dir=prompts_dir,
        summaries_dir=summaries_dir,
    )

    class DummyGenerator:
        def __init__(self, _cfg):
            pytest.fail("Generator should not run while a split-session part is still in progress.")

    monkeypatch.setattr(postprocess_mod, "_GoogleTextGenerator", DummyGenerator)

    with pytest.raises(postprocess_mod.SplitSessionPendingError):
        postprocess_mod.run_postprocess_for_transcript(part_one, config)


def test_run_postprocess_for_split_transcript_stitches_parts(monkeypatch, tmp_path):
    prompts_dir = tmp_path / "prompts"
    summaries_dir = tmp_path / "summaries"
    transcript_root = tmp_path / "transcripts"
    part_one = transcript_root / "Session 28 1_2" / "Session 28 1_2.txt"
    part_two = transcript_root / "Session 28 2_2" / "Session 28 2_2.txt"
    part_one.parent.mkdir(parents=True, exist_ok=True)
    part_two.parent.mkdir(parents=True, exist_ok=True)
    part_one.write_text("Speaker 00 00:00:00 Part one\n", encoding="utf-8")
    part_two.write_text("Speaker 00 00:05:00 Part two\n", encoding="utf-8")
    _write_prompt_files(prompts_dir)
    _write_transcription_marker(part_one, "completed")
    _write_transcription_marker(part_two, "completed")

    config = postprocess_mod.PostProcessConfig(
        enabled=True,
        provider="google",
        model="test-model",
        prompts_dir=prompts_dir,
        summaries_dir=summaries_dir,
    )

    class DummyGenerator:
        calls: list[tuple[str, str | None]] = []

        def __init__(self, cfg):
            assert cfg is config

        def generate(self, *, prompt: str, system_instruction: str | None = None) -> str:
            self.calls.append((prompt, system_instruction))
            if len(self.calls) == 1:
                assert "Part one" in prompt
                assert "Part two" in prompt
                assert prompt.index("Part one") < prompt.index("Part two")
                return "# Analysis\n\nSplit session stitched."
            if len(self.calls) == 2:
                return "<campaign_overview># Campaign Overview\n\nUpdated.</campaign_overview>"
            return "<Summary># Summary\n\nCombined session.</Summary>"

    monkeypatch.setattr(postprocess_mod, "_GoogleTextGenerator", DummyGenerator)

    result = postprocess_mod.run_postprocess_for_transcript(part_two, config)

    assert result.paths.session_folder == summaries_dir / "Session 28"
    assert result.paths.summary_txt == (
        summaries_dir / "Session 28" / "raw_txt" / "session_28_summary_output.txt"
    )
    marker = json.loads(result.paths.completion_marker.read_text(encoding="utf-8"))
    assert marker["source_transcript_paths"] == [str(part_one), str(part_two)]
    assert marker["split_session"] == {"part_total": 2, "source_transcript_count": 2}
