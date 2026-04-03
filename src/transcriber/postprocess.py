from __future__ import annotations

import json
import logging
import os
import re
import shlex
import shutil
import sqlite3
import subprocess
import tempfile
import time
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Callable, Mapping, Optional

logger = logging.getLogger(__name__)

_SPLIT_SESSION_PATTERN = re.compile(
    r"session[\s_-]*(?P<session>\d+)[\s_-]+(?P<part>\d+)_(?P<total>\d+)",
    re.IGNORECASE,
)
_SESSION_PATTERN = re.compile(r"session[\s_-]*(?P<session>\d+)", re.IGNORECASE)
_PLAIN_NUMBER_PATTERN = re.compile(r"(?<!\d)(\d+)(?!\d)")

_PROMPT_FILES = {
    "analysis_system_prompt": "analysis_system_prompt.txt",
    "analysis_prompt_template": "analysis_prompt_template.txt",
    "avarias_background": "avarias_background.txt",
    "base_campaign_overview": "campaign_overview.txt",
    "session_analysis_instructions": "session_analysis_instructions.txt",
    "analysis_example": "analysis_example.txt",
    "campaign_overview_prompt_template": "co_prompt_template.txt",
    "summary_system_prompt": "summary_system_prompt.txt",
    "summary_prompt_template": "summary_prompt_template.txt",
    "summary_instructions": "summary_instructions.txt",
}


@dataclass(frozen=True)
class PostProcessConfig:
    enabled: bool
    provider: str
    model: str
    prompts_dir: Path
    summaries_dir: Path
    api_key_env: str = "GOOGLE_API_KEY"
    raw_text_subdir: str = "raw_txt"
    calls_per_minute: int = 5
    skip_existing: bool = True
    thinking_level: str | None = None
    summary_ready: "SummaryReadyConfig | None" = None

    @property
    def delay_seconds(self) -> float:
        if int(self.calls_per_minute) <= 0:
            return 0.0
        return 60.0 / int(self.calls_per_minute)


@dataclass(frozen=True)
class PostProcessPaths:
    transcript_path: Path
    session_number: int
    session_folder: Path
    raw_text_folder: Path
    analysis_txt: Path
    analysis_docx: Path
    campaign_overview_txt: Path
    campaign_overview_docx: Path
    summary_txt: Path
    summary_docx: Path
    completion_marker: Path


@dataclass(frozen=True)
class SessionIdentity:
    session_number: int
    part_number: int | None = None
    part_total: int | None = None

    @property
    def is_split(self) -> bool:
        return self.part_number is not None and self.part_total is not None


@dataclass(frozen=True)
class TranscriptBundle:
    identity: SessionIdentity
    transcript_path: Path
    source_transcript_paths: tuple[Path, ...]
    missing_parts: tuple[int, ...] = ()

    @property
    def is_complete(self) -> bool:
        return not self.missing_parts


@dataclass(frozen=True)
class PostProcessResult:
    paths: PostProcessPaths
    created_marker: bool


@dataclass(frozen=True)
class SummaryReadyConfig:
    state_target: str
    ssh_options: str = ""
    scenes_count: int = 2


class SplitSessionPendingError(RuntimeError):
    """Raised when a split-session transcript is waiting for remaining parts."""


class UnsupportedTranscriptPostprocessError(RuntimeError):
    """Raised when a transcript name cannot be mapped to a campaign session."""


class _GoogleTextGenerator:
    def __init__(self, config: PostProcessConfig) -> None:
        try:
            from google import genai
            from google.genai import types
        except ImportError as exc:  # pragma: no cover - guarded by dependency install
            raise RuntimeError(
                "Google post-processing requires the google-genai package. Run `uv sync`."
            ) from exc

        api_key = os.getenv(config.api_key_env)
        if not api_key:
            raise RuntimeError(
                f"Missing API key for post-processing. Set ${config.api_key_env} in the environment."
            )

        self._client = genai.Client(api_key=api_key)
        self._types = types
        self._config = config
        self._last_completed_at: Optional[float] = None

    def generate(self, *, prompt: str, system_instruction: Optional[str] = None) -> str:
        if self._last_completed_at is not None:
            elapsed = time.monotonic() - self._last_completed_at
            wait_for = self._config.delay_seconds - elapsed
            if wait_for > 0:
                time.sleep(wait_for)

        config_kwargs = {"response_mime_type": "text/plain"}
        if system_instruction:
            config_kwargs["system_instruction"] = [system_instruction]
        if self._config.thinking_level and self._config.model.lower().startswith("gemini-3"):
            config_kwargs["thinking_config"] = self._types.ThinkingConfig(
                thinking_level=self._config.thinking_level
            )

        response = self._client.models.generate_content(
            model=self._config.model,
            contents=[
                self._types.Content(
                    role="user",
                    parts=[self._types.Part.from_text(text=prompt)],
                )
            ],
            config=self._types.GenerateContentConfig(**config_kwargs),
        )
        self._last_completed_at = time.monotonic()

        text = (getattr(response, "text", None) or "").strip()
        if text:
            return text

        feedback = getattr(response, "prompt_feedback", None)
        block_reason = getattr(feedback, "block_reason", None)
        if block_reason:
            raise RuntimeError(f"Google post-processing blocked the request: {block_reason}")
        raise RuntimeError("Google post-processing returned an empty response.")


def resolve_postprocess_config(cfg: Mapping[str, object] | None) -> Optional[PostProcessConfig]:
    if not isinstance(cfg, Mapping):
        return None

    raw_cfg = cfg.get("postprocess") or {}
    if not isinstance(raw_cfg, Mapping):
        return None
    if not bool(raw_cfg.get("enabled")):
        return None

    prompts_dir = str(raw_cfg.get("prompts_dir") or "").strip()
    summaries_dir = str(raw_cfg.get("summaries_dir") or "").strip()
    model = str(raw_cfg.get("model") or "").strip()
    provider = str(raw_cfg.get("provider") or "google").strip().lower() or "google"
    api_key_env = str(raw_cfg.get("api_key_env") or "GOOGLE_API_KEY").strip() or "GOOGLE_API_KEY"
    raw_text_subdir = str(raw_cfg.get("raw_text_subdir") or "raw_txt").strip() or "raw_txt"
    calls_per_minute_raw = raw_cfg.get("calls_per_minute")
    calls_per_minute = 5 if calls_per_minute_raw is None else int(calls_per_minute_raw)
    skip_existing = bool(raw_cfg.get("skip_existing", True))
    thinking_level_raw = str(raw_cfg.get("thinking_level") or "").strip().lower()
    thinking_level = thinking_level_raw or None
    if thinking_level not in (None, "minimal", "low", "medium", "high"):
        raise SystemExit("postprocess.thinking_level must be one of: minimal, low, medium, high.")
    raw_summary_ready_cfg = raw_cfg.get("summary_ready") or {}
    summary_ready: SummaryReadyConfig | None = None
    if raw_summary_ready_cfg:
        if not isinstance(raw_summary_ready_cfg, Mapping):
            raise SystemExit("postprocess.summary_ready must be a mapping when provided.")
        if bool(raw_summary_ready_cfg.get("enabled")):
            state_target = str(raw_summary_ready_cfg.get("state_target") or "").strip()
            if not state_target:
                raise SystemExit(
                    "postprocess.summary_ready.state_target is required when summary_ready is enabled."
                )
            ssh_options = str(raw_summary_ready_cfg.get("ssh_options") or "").strip()
            scenes_count = max(0, int(raw_summary_ready_cfg.get("scenes_count") or 2))
            summary_ready = SummaryReadyConfig(
                state_target=state_target,
                ssh_options=ssh_options,
                scenes_count=scenes_count,
            )

    if provider != "google":
        raise SystemExit(f"Unsupported postprocess.provider={provider!r}; expected 'google'.")
    if not prompts_dir:
        raise SystemExit("postprocess.prompts_dir is required when post-processing is enabled.")
    if not summaries_dir:
        raise SystemExit("postprocess.summaries_dir is required when post-processing is enabled.")
    if not model:
        raise SystemExit("postprocess.model is required when post-processing is enabled.")

    return PostProcessConfig(
        enabled=True,
        provider=provider,
        model=model,
        prompts_dir=Path(prompts_dir).expanduser(),
        summaries_dir=Path(summaries_dir).expanduser(),
        api_key_env=api_key_env,
        raw_text_subdir=raw_text_subdir,
        calls_per_minute=max(0, calls_per_minute),
        skip_existing=skip_existing,
        thinking_level=thinking_level,
        summary_ready=summary_ready,
    )


def expected_completion_marker_path(
    transcript_path: str | Path,
    config: PostProcessConfig,
) -> Path:
    if not can_postprocess_transcript(transcript_path):
        raise UnsupportedTranscriptPostprocessError(
            f"Transcript does not map to a campaign session: {Path(transcript_path).expanduser()}"
        )
    return build_postprocess_paths(transcript_path, config).completion_marker


def build_postprocess_paths(
    transcript_path: str | Path,
    config: PostProcessConfig,
) -> PostProcessPaths:
    transcript = Path(transcript_path).expanduser()
    session_number = parse_session_identity(transcript).session_number
    session_prefix = f"session_{session_number}"
    session_folder = config.summaries_dir / f"Session {session_number}"
    raw_text_folder = session_folder / config.raw_text_subdir
    return PostProcessPaths(
        transcript_path=transcript,
        session_number=session_number,
        session_folder=session_folder,
        raw_text_folder=raw_text_folder,
        analysis_txt=raw_text_folder / f"{session_prefix}_analysis_output.txt",
        analysis_docx=session_folder / f"{session_prefix}_analysis_output.docx",
        campaign_overview_txt=raw_text_folder / f"{session_prefix}_campaign_overview.txt",
        campaign_overview_docx=session_folder / f"{session_prefix}_campaign_overview.docx",
        summary_txt=raw_text_folder / f"{session_prefix}_summary_output.txt",
        summary_docx=session_folder / f"{session_prefix}_summary_output.docx",
        completion_marker=session_folder / f"{session_prefix}.postprocess.json",
    )


def summary_ready_marker_path(
    transcript_path: str | Path,
    config: PostProcessConfig,
) -> Path:
    """Return the local idempotency marker path for summary-ready publication."""
    paths = build_postprocess_paths(transcript_path, config)
    return _summary_ready_marker_path_for_paths(paths)


def summary_ready_receipt_current(
    transcript_path: str | Path,
    config: PostProcessConfig,
) -> bool:
    """Return true when the summary-ready publication marker is current for a transcript."""
    if config.summary_ready is None:
        return True
    paths = build_postprocess_paths(transcript_path, config)
    bundle = resolve_transcript_bundle(paths.transcript_path)
    if not bundle.is_complete:
        return False
    return _summary_ready_marker_is_current(paths, bundle, config.summary_ready)


def _summary_ready_marker_path_for_paths(paths: PostProcessPaths) -> Path:
    return paths.session_folder / f"session_{paths.session_number}.summary-ready.json"


def run_postprocess_for_transcript(
    transcript_path: str | Path,
    config: PostProcessConfig,
) -> PostProcessResult:
    if not can_postprocess_transcript(transcript_path):
        raise UnsupportedTranscriptPostprocessError(
            f"Transcript does not map to a campaign session: {Path(transcript_path).expanduser()}"
        )
    paths = build_postprocess_paths(transcript_path, config)
    if not paths.transcript_path.exists():
        raise FileNotFoundError(
            f"Transcript not found for post-processing: {paths.transcript_path}"
        )

    transcript_bundle = resolve_transcript_bundle(paths.transcript_path)
    if not transcript_bundle.is_complete:
        missing = ", ".join(str(part) for part in transcript_bundle.missing_parts)
        raise SplitSessionPendingError(
            f"Split-session transcript for session {paths.session_number} is waiting for part(s): "
            f"{missing}"
        )
    can_reuse_existing_outputs = config.skip_existing and _postprocess_outputs_are_current(
        paths,
        transcript_bundle,
    )
    if can_reuse_existing_outputs:
        logger.info("Post-process already complete for session %s", paths.session_number)
        _publish_summary_ready_receipt(
            paths,
            transcript_bundle,
            summary_text=_read_text(paths.summary_txt),
            config=config,
        )
        return PostProcessResult(paths=paths, created_marker=False)

    prompts = _load_prompt_texts(config.prompts_dir)
    previous_campaign_overview = _resolve_previous_campaign_overview(paths, config, prompts)
    transcript_text = _read_transcript_bundle(transcript_bundle)
    client = _GoogleTextGenerator(config)

    analysis_text, _ = _materialize_text_and_docx(
        text_path=paths.analysis_txt,
        docx_path=paths.analysis_docx,
        skip_existing=can_reuse_existing_outputs,
        generate_text=lambda: _generate_analysis(
            client=client,
            prompts=prompts,
            transcript_text=transcript_text,
            previous_campaign_overview=previous_campaign_overview,
        ),
    )
    campaign_overview_text, _ = _materialize_text_and_docx(
        text_path=paths.campaign_overview_txt,
        docx_path=paths.campaign_overview_docx,
        skip_existing=can_reuse_existing_outputs,
        generate_text=lambda: _generate_campaign_overview(
            client=client,
            prompts=prompts,
            previous_campaign_overview=previous_campaign_overview,
            analysis_text=analysis_text,
        ),
    )
    summary_text, _ = _materialize_text_and_docx(
        text_path=paths.summary_txt,
        docx_path=paths.summary_docx,
        skip_existing=can_reuse_existing_outputs,
        generate_text=lambda: _generate_summary(
            client=client,
            prompts=prompts,
            transcript_text=transcript_text,
            analysis_text=analysis_text,
            campaign_overview_text=campaign_overview_text,
        ),
    )

    metadata = {
        "status": "completed",
        "provider": config.provider,
        "model": config.model,
        "session_number": paths.session_number,
        "transcript_path": str(paths.transcript_path),
        "source_transcript_paths": [
            str(path) for path in transcript_bundle.source_transcript_paths
        ],
        "previous_campaign_overview_path": str(
            _resolve_previous_campaign_overview_path(paths, config)
        ),
        "generated_at": time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime()),
        "outputs": {
            "analysis_txt": str(paths.analysis_txt),
            "analysis_docx": str(paths.analysis_docx),
            "campaign_overview_txt": str(paths.campaign_overview_txt),
            "campaign_overview_docx": str(paths.campaign_overview_docx),
            "summary_txt": str(paths.summary_txt),
            "summary_docx": str(paths.summary_docx),
        },
        "summary_preview": summary_text[:200],
    }
    if transcript_bundle.identity.is_split:
        metadata["split_session"] = {
            "part_total": transcript_bundle.identity.part_total,
            "source_transcript_count": len(transcript_bundle.source_transcript_paths),
        }
    _atomic_write_text(paths.completion_marker, json.dumps(metadata, indent=2) + "\n")
    logger.warning(
        "Post-process complete for session %s (summary=%s)",
        paths.session_number,
        paths.summary_docx,
    )
    _publish_summary_ready_receipt(
        paths,
        transcript_bundle,
        summary_text=summary_text,
        config=config,
        client=client,
    )
    return PostProcessResult(paths=paths, created_marker=True)


def save_markdown_as_docx(markdown_text: str, docx_path: str | Path) -> None:
    try:
        import markdown as markdown_lib
        from bs4 import BeautifulSoup, NavigableString, Tag
        from docx import Document
        from docx.oxml.ns import qn
    except ImportError as exc:  # pragma: no cover - guarded by dependency install
        raise RuntimeError(
            "DOCX export requires markdown, beautifulsoup4, and python-docx. Run `uv sync`."
        ) from exc

    def set_run_mono(run) -> None:
        run.font.name = "Courier New"
        try:
            r_fonts = run._element.rPr.rFonts
            r_fonts.set(qn("w:eastAsia"), "Courier New")
            r_fonts.set(qn("w:ascii"), "Courier New")
            r_fonts.set(qn("w:hAnsi"), "Courier New")
        except Exception:
            pass

    def add_inline(paragraph, node) -> None:
        if isinstance(node, NavigableString):
            text = str(node)
            if text:
                paragraph.add_run(text)
            return
        if not isinstance(node, Tag):
            return

        name = node.name.lower()
        if name in {"strong", "b"}:
            run = paragraph.add_run(node.get_text())
            run.bold = True
            return
        if name in {"em", "i"}:
            run = paragraph.add_run(node.get_text())
            run.italic = True
            return
        if name == "code":
            run = paragraph.add_run(node.get_text())
            set_run_mono(run)
            return
        for child in node.children:
            add_inline(paragraph, child)

    def add_paragraph(document, paragraph_tag) -> None:
        paragraph = document.add_paragraph()
        for child in paragraph_tag.children:
            add_inline(paragraph, child)

    def add_list(document, list_tag, *, ordered: bool) -> None:
        style_name = "List Number" if ordered else "List Bullet"
        for item in list_tag.find_all("li", recursive=False):
            paragraph = document.add_paragraph(style=style_name)
            for child in item.children:
                if isinstance(child, Tag) and child.name.lower() in {"ul", "ol"}:
                    continue
                add_inline(paragraph, child)
            for nested in item.find_all(["ul", "ol"], recursive=False):
                add_list(document, nested, ordered=nested.name.lower() == "ol")

    def add_code_block(document, text: str) -> None:
        for line in text.splitlines() or [""]:
            paragraph = document.add_paragraph()
            run = paragraph.add_run(line)
            set_run_mono(run)

    def add_table(document, table_tag) -> None:
        rows = table_tag.find_all("tr")
        if not rows:
            return
        first_row = rows[0].find_all(["th", "td"])
        table = document.add_table(rows=len(rows), cols=len(first_row))
        for row_index, row in enumerate(rows):
            cells = row.find_all(["th", "td"])
            for col_index, cell in enumerate(cells):
                table.cell(row_index, col_index).text = cell.get_text()

    html = markdown_lib.markdown(
        markdown_text,
        extensions=["fenced_code", "tables", "sane_lists"],
    )
    soup = BeautifulSoup(html, "html.parser")
    document = Document()

    for element in soup.contents:
        if isinstance(element, NavigableString):
            if str(element).strip():
                document.add_paragraph(str(element).strip())
            continue
        if not isinstance(element, Tag):
            continue

        name = element.name.lower()
        if name in {"h1", "h2", "h3", "h4"}:
            document.add_heading(element.get_text(), level=int(name[1]))
        elif name == "p":
            add_paragraph(document, element)
        elif name == "ul":
            add_list(document, element, ordered=False)
        elif name == "ol":
            add_list(document, element, ordered=True)
        elif name == "blockquote":
            try:
                document.add_paragraph(element.get_text(), style="Intense Quote")
            except Exception:
                paragraph = document.add_paragraph()
                run = paragraph.add_run(element.get_text())
                run.italic = True
        elif name == "pre":
            code_node = element.find("code")
            add_code_block(document, code_node.get_text() if code_node else element.get_text())
        elif name == "table":
            add_table(document, element)
        else:
            document.add_paragraph(element.get_text())

    _atomic_save_docx(document, Path(docx_path).expanduser())


def parse_session_identity(path: str | Path) -> SessionIdentity:
    transcript_path = Path(path).expanduser()
    for candidate in _session_identity_candidates(transcript_path):
        split_match = _SPLIT_SESSION_PATTERN.search(candidate)
        if split_match:
            return SessionIdentity(
                session_number=int(split_match.group("session")),
                part_number=int(split_match.group("part")),
                part_total=int(split_match.group("total")),
            )
    for candidate in _session_identity_candidates(transcript_path):
        session_match = _SESSION_PATTERN.search(candidate)
        if session_match:
            return SessionIdentity(session_number=int(session_match.group("session")))
    for candidate in _session_identity_candidates(transcript_path):
        number_match = _PLAIN_NUMBER_PATTERN.search(candidate)
        if number_match:
            return SessionIdentity(session_number=int(number_match.group(1)))
    raise ValueError(f"Could not infer a session number from transcript path: {transcript_path}")


def has_explicit_session_identity(path: str | Path) -> bool:
    transcript_path = Path(path).expanduser()
    for candidate in _session_identity_candidates(transcript_path):
        if _SPLIT_SESSION_PATTERN.search(candidate):
            return True
    for candidate in _session_identity_candidates(transcript_path):
        if _SESSION_PATTERN.search(candidate):
            return True
    return False


def can_postprocess_transcript(path: str | Path) -> bool:
    return has_explicit_session_identity(path)


def resolve_transcript_bundle(path: str | Path) -> TranscriptBundle:
    transcript_path = Path(path).expanduser()
    identity = parse_session_identity(transcript_path)
    if not identity.is_split:
        return TranscriptBundle(
            identity=identity,
            transcript_path=transcript_path,
            source_transcript_paths=(transcript_path,),
        )

    output_root = transcript_path.parent.parent
    by_part: dict[int, Path] = {}
    for candidate in sorted(output_root.rglob("*.txt")):
        try:
            candidate_identity = parse_session_identity(candidate)
        except ValueError:
            continue
        if candidate_identity.session_number != identity.session_number:
            continue
        if not candidate_identity.is_split:
            continue
        if candidate_identity.part_total != identity.part_total:
            continue
        if _transcription_marker_status(candidate) not in (None, "completed"):
            continue
        assert candidate_identity.part_number is not None
        by_part.setdefault(candidate_identity.part_number, candidate)

    source_transcript_paths: list[Path] = []
    missing_parts: list[int] = []
    assert identity.part_total is not None
    for part_number in range(1, identity.part_total + 1):
        candidate = by_part.get(part_number)
        if candidate is None:
            missing_parts.append(part_number)
            continue
        source_transcript_paths.append(candidate)

    return TranscriptBundle(
        identity=identity,
        transcript_path=transcript_path,
        source_transcript_paths=tuple(source_transcript_paths),
        missing_parts=tuple(missing_parts),
    )


def split_session_ready_for_postprocess(path: str | Path) -> bool:
    return resolve_transcript_bundle(path).is_complete


def _session_identity_candidates(path: Path) -> tuple[str, ...]:
    grandparent_name = path.parent.parent.name if path.parent.parent != path.parent else ""
    return (
        path.stem,
        path.name,
        path.parent.name,
        grandparent_name,
    )


def _load_prompt_texts(prompts_dir: Path) -> dict[str, str]:
    loaded: dict[str, str] = {}
    for key, filename in _PROMPT_FILES.items():
        prompt_path = prompts_dir / filename
        loaded[key] = _read_text(prompt_path)
    return loaded


def _resolve_previous_campaign_overview(
    paths: PostProcessPaths,
    config: PostProcessConfig,
    prompts: Mapping[str, str],
) -> str:
    previous_path = _resolve_previous_campaign_overview_path(paths, config)
    if previous_path.exists():
        return _read_text(previous_path)
    logger.info(
        "Previous campaign overview missing for session %s; using base prompt file.",
        paths.session_number,
    )
    return prompts["base_campaign_overview"]


def _resolve_previous_campaign_overview_path(
    paths: PostProcessPaths,
    config: PostProcessConfig,
) -> Path:
    if paths.session_number <= 1:
        return config.prompts_dir / _PROMPT_FILES["base_campaign_overview"]
    previous_session = paths.session_number - 1
    return (
        config.summaries_dir
        / f"Session {previous_session}"
        / config.raw_text_subdir
        / f"session_{previous_session}_campaign_overview.txt"
    )


def _generate_analysis(
    *,
    client: _GoogleTextGenerator,
    prompts: Mapping[str, str],
    transcript_text: str,
    previous_campaign_overview: str,
) -> str:
    prompt = prompts["analysis_prompt_template"].format(
        avarias_bg=prompts["avarias_background"],
        campaign_ovw=previous_campaign_overview,
        session_trans=transcript_text,
        session_analysis_inst=prompts["session_analysis_instructions"],
        analysis_example=prompts["analysis_example"],
    )
    return client.generate(
        prompt=prompt,
        system_instruction=prompts["analysis_system_prompt"],
    )


def _generate_campaign_overview(
    *,
    client: _GoogleTextGenerator,
    prompts: Mapping[str, str],
    previous_campaign_overview: str,
    analysis_text: str,
) -> str:
    prompt = prompts["campaign_overview_prompt_template"].format(
        previous_campaign_overview=previous_campaign_overview,
        session_analysis=analysis_text,
    )
    response_text = client.generate(prompt=prompt)
    extracted = _extract_tagged_text(response_text, "campaign_overview")
    if extracted:
        return extracted
    logger.warning(
        "Campaign overview response was missing <campaign_overview> tags; using full text."
    )
    return response_text


def _generate_summary(
    *,
    client: _GoogleTextGenerator,
    prompts: Mapping[str, str],
    transcript_text: str,
    analysis_text: str,
    campaign_overview_text: str,
) -> str:
    prompt = prompts["summary_prompt_template"].format(
        avarias_bg=prompts["avarias_background"],
        campaign_ovw=campaign_overview_text,
        session_trans=transcript_text,
        session_analysis=analysis_text,
        summary_inst=prompts["summary_instructions"],
    )
    response_text = client.generate(
        prompt=prompt,
        system_instruction=prompts["summary_system_prompt"],
    )
    extracted = _extract_tagged_text(response_text, "summary")
    if extracted:
        return extracted
    logger.warning("Summary response was missing <Summary> tags; using full text.")
    return response_text


def _extract_tagged_text(text: str, tag_name: str) -> str:
    pattern = re.compile(
        rf"<{re.escape(tag_name)}>(.*?)</{re.escape(tag_name)}>",
        re.IGNORECASE | re.DOTALL,
    )
    match = pattern.search(text)
    return match.group(1).strip() if match else ""


def _extract_tagged_blocks(text: str, tag_name: str) -> list[str]:
    pattern = re.compile(
        rf"<{re.escape(tag_name)}>(.*?)</{re.escape(tag_name)}>",
        re.IGNORECASE | re.DOTALL,
    )
    return [match.strip() for match in pattern.findall(text) if match.strip()]


def _build_summary_excerpt(text: str, limit: int = 280) -> str:
    excerpt = " ".join(text.split())
    if limit <= 0 or len(excerpt) <= limit:
        return excerpt
    return excerpt[: max(limit - 3, 0)].rstrip() + "..."


def _fallback_announcement_scenes(summary_text: str, count: int) -> list[str]:
    if count <= 0:
        return []
    candidates: list[str] = []
    for block in re.split(r"\n\s*\n", summary_text):
        line = " ".join(block.split()).strip()
        if not line:
            continue
        if line.startswith("#"):
            continue
        candidates.append(line)
    if not candidates:
        fallback = _build_summary_excerpt(summary_text, 280)
        return [fallback] if fallback else []
    return candidates[:count]


def _generate_announcement_scenes(
    client: _GoogleTextGenerator,
    *,
    summary_text: str,
    scenes_count: int,
) -> list[str]:
    if scenes_count <= 0:
        return []
    prompt = (
        "Select the funniest or most impactful scenes from this tabletop RPG session summary.\n"
        f"Return exactly {scenes_count} short, vivid scene snippets.\n"
        "Use only this XML format and no extra text:\n"
        "<announcement_scenes><scene>...</scene><scene>...</scene></announcement_scenes>\n\n"
        f"{summary_text.strip()}"
    )
    response_text = client.generate(
        prompt=prompt,
        system_instruction=(
            "You write short Discord-ready scene callouts from RPG session summaries. "
            "Favor funny, surprising, or emotionally strong moments. Return only XML."
        ),
    )
    scenes = [
        " ".join(scene.split())
        for scene in _extract_tagged_blocks(response_text, "scene")
        if scene.strip()
    ]
    if scenes:
        return scenes[:scenes_count]
    logger.warning("Announcement scene response was missing <scene> tags; using summary fallback.")
    return _fallback_announcement_scenes(summary_text, scenes_count)


def _materialize_text_and_docx(
    *,
    text_path: Path,
    docx_path: Path,
    skip_existing: bool,
    generate_text: Callable[[], str],
) -> tuple[str, bool]:
    if skip_existing and text_path.exists():
        text = _read_text(text_path)
        if not docx_path.exists():
            save_markdown_as_docx(text, docx_path)
        return text, False

    text = generate_text().strip()
    if not text:
        raise RuntimeError(f"Refusing to write empty post-process output: {text_path}")
    _atomic_write_text(text_path, text + "\n")
    save_markdown_as_docx(text, docx_path)
    return text, True


def _summary_ready_marker_is_current(
    paths: PostProcessPaths,
    bundle: TranscriptBundle,
    summary_ready_config: SummaryReadyConfig,
) -> bool:
    del summary_ready_config
    marker_path = _summary_ready_marker_path_for_paths(paths)
    metadata = _load_postprocess_metadata(marker_path)
    if metadata is None:
        return False
    if not str(metadata.get("document_url") or "").strip():
        return False
    expected_sources = [str(path) for path in bundle.source_transcript_paths]
    actual_sources = metadata.get("source_transcript_paths")
    if not isinstance(actual_sources, list):
        return False
    if metadata.get("transcript_path") != str(paths.transcript_path):
        return False
    if [str(item) for item in actual_sources] != expected_sources:
        return False
    try:
        newest_source_mtime = max(
            path.stat().st_mtime
            for path in (
                *bundle.source_transcript_paths,
                paths.summary_txt,
                paths.completion_marker,
            )
        )
        marker_mtime = marker_path.stat().st_mtime
    except OSError:
        return False
    return marker_mtime >= newest_source_mtime


def _postprocess_complete(paths: PostProcessPaths) -> bool:
    return all(path.exists() for path in _required_postprocess_files(paths))


def _required_postprocess_files(paths: PostProcessPaths) -> tuple[Path, ...]:
    return (
        paths.analysis_txt,
        paths.analysis_docx,
        paths.campaign_overview_txt,
        paths.campaign_overview_docx,
        paths.summary_txt,
        paths.summary_docx,
        paths.completion_marker,
    )


def _load_postprocess_metadata(marker_path: Path) -> dict[str, object] | None:
    if not marker_path.exists():
        return None
    try:
        payload = json.loads(marker_path.read_text(encoding="utf-8"))
    except Exception:
        return None
    return payload if isinstance(payload, dict) else None


def _postprocess_outputs_are_current(
    paths: PostProcessPaths,
    bundle: TranscriptBundle,
) -> bool:
    required_files = _required_postprocess_files(paths)
    if not all(path.exists() for path in required_files):
        return False

    metadata = _load_postprocess_metadata(paths.completion_marker)
    if metadata is None:
        return False

    expected_sources = [str(path) for path in bundle.source_transcript_paths]
    actual_sources = metadata.get("source_transcript_paths")
    if not isinstance(actual_sources, list):
        return False
    if metadata.get("transcript_path") != str(paths.transcript_path):
        return False
    if [str(item) for item in actual_sources] != expected_sources:
        return False

    try:
        newest_source_mtime = max(path.stat().st_mtime for path in bundle.source_transcript_paths)
        oldest_output_mtime = min(path.stat().st_mtime for path in required_files)
    except OSError:
        return False
    return oldest_output_mtime >= newest_source_mtime


def _transcription_marker_payload(transcript_path: str | Path) -> dict[str, Any] | None:
    marker_path = _transcription_completion_marker_path(transcript_path)
    if not marker_path.exists():
        return None
    try:
        payload = json.loads(marker_path.read_text(encoding="utf-8"))
    except Exception:
        return None
    return payload if isinstance(payload, dict) else None


def _load_json_payload(path: Path) -> dict[str, Any] | None:
    try:
        payload = json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return None
    return payload if isinstance(payload, dict) else None


def _load_manifest_from_marker(marker_payload: dict[str, Any]) -> dict[str, Any] | None:
    candidates: list[Path] = []
    manifest_path_raw = str(marker_payload.get("input_manifest_path") or "").strip()
    if manifest_path_raw:
        candidates.append(Path(manifest_path_raw).expanduser())

    input_path_raw = str(marker_payload.get("input_path") or "").strip()
    if input_path_raw:
        input_path = Path(input_path_raw).expanduser()
        candidates.append(input_path.with_suffix(".json"))
        candidates.append(input_path.parent / "processed" / f"{input_path.stem}.json")

    for candidate in candidates:
        if not candidate.exists():
            continue
        payload = _load_json_payload(candidate)
        if payload is not None:
            return payload
    return None


def _summary_ready_context(
    paths: PostProcessPaths,
) -> tuple[str, str] | None:
    marker_payload = _transcription_marker_payload(paths.transcript_path) or {}
    manifest_payload = _load_manifest_from_marker(marker_payload) or {}
    recording_id = str(
        marker_payload.get("recording_id") or manifest_payload.get("recording_id") or ""
    ).strip()
    if not recording_id:
        return None
    title = str(
        marker_payload.get("title")
        or manifest_payload.get("preferred_basename")
        or manifest_payload.get("requested_name")
        or f"Session {paths.session_number}"
    ).strip()
    return recording_id, title or f"Session {paths.session_number}"


def _drivefs_metadata_db_paths() -> list[Path]:
    roots = Path("/mnt/c/Users")
    try:
        if not roots.exists():
            return []
        user_dirs = list(roots.iterdir())
    except OSError:
        return []
    candidates: list[Path] = []
    for user_dir in user_dirs:
        drivefs_root = user_dir / "AppData" / "Local" / "Google" / "DriveFS"
        try:
            if not drivefs_root.exists():
                continue
            account_dirs = list(drivefs_root.iterdir())
        except OSError:
            continue
        for account_dir in account_dirs:
            candidate = account_dir / "mirror_metadata_sqlite.db"
            try:
                if candidate.is_file():
                    candidates.append(candidate)
            except OSError:
                continue
    return sorted(candidates)


def _drivefs_relative_parts(path: Path) -> tuple[str, ...] | None:
    normalized_parts = path.expanduser().parts
    for anchor in ("My Drive", "Shared drives"):
        if anchor in normalized_parts:
            start = normalized_parts.index(anchor)
            return normalized_parts[start:]
    return None


def _copy_db_snapshot(path: Path) -> Path:
    snapshot = Path(tempfile.mkstemp(prefix=f"{path.stem}.", suffix=".sqlite")[1])
    shutil.copy2(path, snapshot)
    return snapshot


def _drivefs_cloud_id_for_parts(db_path: Path, parts: tuple[str, ...]) -> str | None:
    if not parts:
        return None
    snapshot_path = _copy_db_snapshot(db_path)
    try:
        conn = sqlite3.connect(str(snapshot_path))
        try:
            candidates = [
                row[0]
                for row in conn.execute(
                    """
                    select stable_id
                    from items
                    where local_title = ? and is_tombstone = 0
                    """,
                    (parts[0],),
                ).fetchall()
            ]
            for part in parts[1:]:
                if not candidates:
                    return None
                next_candidates: list[int] = []
                for parent_stable_id in candidates:
                    next_candidates.extend(
                        row[0]
                        for row in conn.execute(
                            """
                            select i.stable_id
                            from items i
                            join stable_parents sp on sp.item_stable_id = i.stable_id
                            where i.local_title = ?
                              and i.is_tombstone = 0
                              and sp.parent_stable_id = ?
                            """,
                            (part, parent_stable_id),
                        ).fetchall()
                    )
                candidates = next_candidates
            if not candidates:
                return None
            row = conn.execute(
                """
                select cloud_id
                from stable_ids
                where stable_id = ?
                """,
                (candidates[0],),
            ).fetchone()
            if row and isinstance(row[0], str) and row[0].strip():
                return row[0].strip()
            return None
        finally:
            conn.close()
    finally:
        snapshot_path.unlink(missing_ok=True)


def _resolve_summary_document_url(summary_docx_path: Path) -> str | None:
    parts = _drivefs_relative_parts(summary_docx_path)
    if not parts:
        return None
    for db_path in _drivefs_metadata_db_paths():
        try:
            cloud_id = _drivefs_cloud_id_for_parts(db_path, parts)
        except Exception as exc:  # noqa: BLE001
            logger.warning("Failed to query DriveFS metadata at %s: %s", db_path, exc)
            continue
        if cloud_id:
            return f"https://drive.google.com/open?id={cloud_id}"
    return None


def _is_remote_target(target: str) -> bool:
    return ":" in target and not target.startswith(("/", ".", "~"))


def _parse_remote_target(target: str) -> tuple[str, str]:
    host, separator, remote_path = target.partition(":")
    if not separator or not host or not remote_path:
        raise ValueError(f"Unsupported remote target: {target!r}")
    return host, remote_path


def _build_ssh_command(
    summary_ready: SummaryReadyConfig, host: str, remote_command: str
) -> list[str]:
    command = ["ssh"]
    if summary_ready.ssh_options:
        command.extend(shlex.split(summary_ready.ssh_options))
    command.extend([host, remote_command])
    return command


def _run_checked(
    command: list[str], *, input_text: str | None = None
) -> subprocess.CompletedProcess[str]:
    return subprocess.run(
        command,
        check=True,
        capture_output=True,
        input=input_text,
        text=True,
    )


def _load_state_receipt(
    summary_ready: SummaryReadyConfig,
    *,
    recording_id: str,
    suffix: str,
) -> dict[str, Any] | None:
    receipt_name = f"{recording_id}.{suffix}.json"
    if _is_remote_target(summary_ready.state_target):
        host, remote_dir = _parse_remote_target(summary_ready.state_target)
        remote_receipt = str(Path(remote_dir) / receipt_name)
        command = _build_ssh_command(
            summary_ready,
            host,
            f"cat {shlex.quote(remote_receipt)}",
        )
        try:
            result = _run_checked(command)
        except subprocess.CalledProcessError:
            return None
        try:
            payload = json.loads(result.stdout)
        except Exception:
            return None
        return payload if isinstance(payload, dict) else None

    receipt_path = Path(summary_ready.state_target).expanduser() / receipt_name
    if not receipt_path.exists():
        return None
    return _load_json_payload(receipt_path)


def _write_summary_ready_receipt(
    summary_ready: SummaryReadyConfig,
    *,
    recording_id: str,
    title: str,
    summary_excerpt: str,
    drive_url: str,
    completed_at: str,
    announcement_scenes: list[str],
    document_url: str = "",
    status: str = "ready",
) -> None:
    receipt_name = f"{recording_id}.summary-ready.json"
    payload: dict[str, Any] = {
        "recording_id": recording_id,
        "title": title,
        "summary_excerpt": summary_excerpt,
        "drive_url": drive_url,
        "completed_at": completed_at,
        "status": status,
    }
    cleaned_scenes = [scene.strip() for scene in announcement_scenes if scene.strip()]
    if cleaned_scenes:
        payload["announcement_scenes"] = cleaned_scenes
    if document_url.strip():
        payload["document_url"] = document_url.strip()
    payload_text = json.dumps(payload, indent=2, sort_keys=True) + "\n"

    if _is_remote_target(summary_ready.state_target):
        host, remote_dir = _parse_remote_target(summary_ready.state_target)
        remote_receipt = str(Path(remote_dir) / receipt_name)
        temp_receipt = f"{remote_receipt}.tmp.{os.getpid()}"
        remote_command = (
            f"mkdir -p {shlex.quote(remote_dir)} && "
            f"cat > {shlex.quote(temp_receipt)} && "
            f"mv {shlex.quote(temp_receipt)} {shlex.quote(remote_receipt)}"
        )
        _run_checked(
            _build_ssh_command(summary_ready, host, remote_command),
            input_text=payload_text,
        )
        return

    state_dir = Path(summary_ready.state_target).expanduser()
    state_dir.mkdir(parents=True, exist_ok=True)
    temp_receipt = state_dir / f".{receipt_name}.tmp.{os.getpid()}"
    temp_receipt.write_text(payload_text, encoding="utf-8")
    os.replace(temp_receipt, state_dir / receipt_name)


def _write_summary_ready_marker(
    paths: PostProcessPaths,
    bundle: TranscriptBundle,
    *,
    recording_id: str,
    drive_url: str,
    document_url: str,
) -> None:
    marker_path = _summary_ready_marker_path_for_paths(paths)
    metadata = {
        "recording_id": recording_id,
        "transcript_path": str(paths.transcript_path),
        "source_transcript_paths": [str(path) for path in bundle.source_transcript_paths],
        "summary_path": str(paths.summary_txt),
        "drive_url": drive_url,
        "document_url": document_url,
        "published_at": time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime()),
    }
    _atomic_write_text(marker_path, json.dumps(metadata, indent=2) + "\n")


def _publish_summary_ready_receipt(
    paths: PostProcessPaths,
    bundle: TranscriptBundle,
    *,
    summary_text: str,
    config: PostProcessConfig,
    client: _GoogleTextGenerator | None = None,
) -> None:
    summary_ready = config.summary_ready
    if summary_ready is None:
        return
    if _summary_ready_marker_is_current(paths, bundle, summary_ready):
        return

    context = _summary_ready_context(paths)
    if context is None:
        logger.warning(
            "Skipping summary-ready receipt for %s: missing Craig recording id",
            paths.transcript_path,
        )
        return
    recording_id, title = context
    drive_receipt = _load_state_receipt(
        summary_ready,
        recording_id=recording_id,
        suffix="drive-uploaded",
    )
    drive_url = str((drive_receipt or {}).get("url") or "").strip()
    if not drive_url:
        logger.warning(
            "Skipping summary-ready receipt for %s: drive-uploaded receipt is not available yet",
            recording_id,
        )
        return
    document_url = _resolve_summary_document_url(paths.summary_docx)
    if not document_url:
        logger.warning(
            "Skipping summary-ready receipt for %s: summary document Drive link is not available yet",
            recording_id,
        )
        return

    generator = client or _GoogleTextGenerator(config)
    announcement_scenes = _generate_announcement_scenes(
        generator,
        summary_text=summary_text,
        scenes_count=summary_ready.scenes_count,
    )
    completion_payload = _load_postprocess_metadata(paths.completion_marker) or {}
    completed_at = str(
        completion_payload.get("generated_at") or time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime())
    ).strip()
    _write_summary_ready_receipt(
        summary_ready,
        recording_id=recording_id,
        title=title,
        summary_excerpt=_build_summary_excerpt(summary_text),
        drive_url=drive_url,
        completed_at=completed_at,
        announcement_scenes=announcement_scenes,
        document_url=document_url,
    )
    _write_summary_ready_marker(
        paths,
        bundle,
        recording_id=recording_id,
        drive_url=drive_url,
        document_url=document_url,
    )


def _read_text(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8")
    except FileNotFoundError:
        raise
    except Exception as exc:  # noqa: BLE001
        raise RuntimeError(f"Failed to read {path}: {exc}") from exc


def _read_transcript_bundle(bundle: TranscriptBundle) -> str:
    chunks = []
    for path in bundle.source_transcript_paths:
        chunk = _read_text(path).strip()
        if chunk:
            chunks.append(chunk)
    combined = "\n\n".join(chunks).strip()
    if not combined:
        raise RuntimeError(
            f"Refusing to run post-processing with an empty transcript bundle for {bundle.transcript_path}"
        )
    return combined


def _transcription_completion_marker_path(transcript_path: str | Path) -> Path:
    transcript = Path(transcript_path).expanduser()
    return transcript.parent / f"{transcript.stem}.transcribe.json"


def _transcription_marker_status(transcript_path: str | Path) -> str | None:
    marker_path = _transcription_completion_marker_path(transcript_path)
    if not marker_path.exists():
        return None
    try:
        payload = json.loads(marker_path.read_text(encoding="utf-8"))
    except Exception:
        return "invalid"
    status = str(payload.get("status") or "").strip().lower()
    return status or "invalid"


def _atomic_write_text(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    tmp_file = tempfile.NamedTemporaryFile(
        "w",
        encoding="utf-8",
        dir=path.parent,
        prefix=f".{path.name}.",
        suffix=".tmp",
        delete=False,
    )
    tmp_path = Path(tmp_file.name)
    try:
        with tmp_file:
            tmp_file.write(text)
        tmp_path.replace(path)
    finally:
        if tmp_path.exists():
            tmp_path.unlink(missing_ok=True)


def _atomic_save_docx(document, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    fd, tmp_name = tempfile.mkstemp(
        dir=path.parent,
        prefix=f".{path.stem}.",
        suffix=".docx.tmp",
    )
    os.close(fd)
    tmp_path = Path(tmp_name)
    try:
        document.save(tmp_path)
        tmp_path.replace(path)
    finally:
        if tmp_path.exists():
            tmp_path.unlink(missing_ok=True)
